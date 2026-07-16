import os
from typing import Optional, Dict, List, Any
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np


class APAWordExporter:
    def __init__(self, title: str = "Statistical Analysis Report"):
        self.doc = Document()
        self._setup_styles()
        self.title = title
        self._add_title_page(title)

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_after = Pt(0)

    def _set_cell_font(self, cell, text: str, bold: bool = False, size: int = 10, italic: bool = False):
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _shade_cell(self, cell, color: str = "D9E2F3"):
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)

    def _add_title_page(self, title: str):
        for _ in range(6):
            self.doc.add_paragraph('')
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(22)
        run.font.name = 'Times New Roman'

        self.doc.add_paragraph('')
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Statistical Analysis Report")
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'

        self.doc.add_page_break()

    def add_heading(self, text: str, level: int = 1):
        h = self.doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
        return h

    def add_paragraph(self, text: str, bold: bool = False, italic: bool = False) -> Any:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic
        p.paragraph_format.line_spacing = 2.0
        return p

    def add_apa_table(self, headers: List[str], rows: List[List[str]],
                      title: str = "", caption: str = "", note: str = "") -> Any:
        if title:
            self.add_paragraph(title, bold=True)

        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            self._set_cell_font(cell, h, bold=True, size=9)
            self._shade_cell(cell)

        for i, row in enumerate(rows):
            for j, val in enumerate(row):
                cell = table.rows[i + 1].cells[j]
                self._set_cell_font(cell, str(val), size=9)

        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"Table. {caption}")
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

        if note:
            p = self.doc.add_paragraph()
            run = p.add_run(f"Note. {note}")
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
            run.italic = True

        self.doc.add_paragraph('')
        return table

    def add_figure(self, fig: plt.Figure, caption: str = "", width: float = 5.5):
        img_path = os.path.join(os.path.dirname(__file__) or '.', '_temp_fig.png')
        fig.savefig(img_path, dpi=300, bbox_inches='tight', facecolor='white')
        self.doc.add_picture(img_path, width=Inches(width))
        os.remove(img_path)

        if caption:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"Figure. {caption}")
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

        self.doc.add_paragraph('')
        plt.close(fig)

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, filepath: str):
        self.doc.save(filepath)
        return filepath
